import pandas as pd
from docx import Document
import os

# Define the source file and read the Excel data from the "DATA" sheet
source = 'StateCommissions_PHA_pulled 6.13.2024_1214pm.xlsx'
df = pd.read_excel(source, sheet_name='DATA')

# Replace NaN with empty strings
df = df.fillna('')

# Define the placeholder replacements
placeholders = {
    '<AUTHORIZED_REP>': 'auth name',
    '<ORGANIZATION>': 'Org',
    '<Org Street Addr 1>': 'Adress1',
    '<Org Street Addr 2>': 'Address 2',
    '<ORG CITY>': 'City',
    '<ORG STATE>': 'State',
    '<ORG ZIP>': 'zip',
    '<Authorized Rep Name>': 'auth name',
}

# Function to replace placeholders in the document
def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            for key, value in replacements.items():
                if key in text:
                    text = text.replace(key, str(value) if value is not None else "")
            run.text = text
            print(f"Paragraph after replacement: {run.text}")  # Debugging line

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        text = run.text
                        for key, value in replacements.items():
                            if key in text:
                                text = text.replace(key, str(value) if value is not None else "")
                        run.text = text
                        print(f"Cell after replacement: {run.text}")  # Debugging line

# Ensure the output directory exists
output_dir = 'outputs'
os.makedirs(output_dir, exist_ok=True)

# Process only the first row in the DataFrame
row = df.iloc[0]
doc = Document('FY_24_PHAComm_Notification__Ltr_TEMPLATE (3).docx')

replacements = {key: row[value] for key, value in placeholders.items()}
print(f"Replacements for {row['File_Name']}: {replacements}")  # Debugging line
replace_placeholders(doc, replacements)

file_name = row['File_Name'] + '.docx'
doc.save(os.path.join(output_dir, file_name))

print("Mail merge completed successfully.")
